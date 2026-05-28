"""PURE document loader.

Reads documents from a directory and extracts raw text with structural
preservation.  Supports: PDF, DOCX, HTML/HTM, TXT, Markdown.

Key principle: preserve line breaks, headings, and paragraph structure as
much as possible — aggressive normalization is deferred to downstream modules
that know what they are looking for.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

import structlog

from reqlens_benchmark_builder.config import get_settings

logger = structlog.get_logger(__name__)

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".html", ".htm", ".txt", ".md"}


# ── Domain object ─────────────────────────────────────────────────────────────

@dataclass
class PureDocument:
    doc_id: str
    path: Path
    file_type: str
    text: str             # raw extracted text, structural cues preserved
    char_count: int = 0

    def __post_init__(self) -> None:
        self.char_count = len(self.text)


# ── Format-specific extractors ────────────────────────────────────────────────

def _extract_pdf(path: Path) -> str:
    """Extract text from PDF, preserving page markers and layout cues."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf is required: pip install pypdf")

    reader = PdfReader(str(path))
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            # Page marker helps section detector notice page boundaries
            pages.append(f"\n[Page {i + 1}]\n{text}")
    full = "\n".join(pages).strip()
    logger.debug("pure_loader.pdf_extracted", path=str(path), chars=len(full))
    return full


def _extract_docx(path: Path) -> str:
    """Extract DOCX text, preserving heading levels as numbered prefixes."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = DocxDocument(str(path))
    blocks: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        # Represent headings with Markdown-style # prefix so section detector works
        if "heading 1" in style:
            blocks.append(f"\n# {text}")
        elif "heading 2" in style:
            blocks.append(f"\n## {text}")
        elif "heading 3" in style:
            blocks.append(f"\n### {text}")
        elif "heading 4" in style:
            blocks.append(f"\n#### {text}")
        else:
            blocks.append(text)

    full = "\n\n".join(blocks).strip()
    logger.debug("pure_loader.docx_extracted", path=str(path), chars=len(full))
    return full


def _extract_html(path: Path) -> str:
    """Extract readable text from HTML, stripping scripts and styles."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("beautifulsoup4 is required: pip install beautifulsoup4 lxml")

    html = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "lxml")

    # Remove noise tags
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    # Preserve heading semantics as Markdown-style markers
    for level in range(1, 7):
        for h in soup.find_all(f"h{level}"):
            prefix = "#" * level
            h.replace_with(f"\n{prefix} {h.get_text().strip()}\n")

    lines = [
        line.strip()
        for line in soup.get_text("\n").splitlines()
        if line.strip()
    ]
    full = "\n".join(lines)
    # Collapse runs of 3+ blank lines
    full = re.sub(r"\n{3,}", "\n\n", full).strip()
    logger.debug("pure_loader.html_extracted", path=str(path), chars=len(full))
    return full


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace").strip()


_EXTRACTORS = {
    ".pdf":  _extract_pdf,
    ".docx": _extract_docx,
    ".html": _extract_html,
    ".htm":  _extract_html,
    ".txt":  _extract_txt,
    ".md":   _extract_txt,
}


# ── Public loader ─────────────────────────────────────────────────────────────

def load_pure_documents() -> list[PureDocument]:
    """Load all supported documents from the configured PURE input directory.

    Documents are sorted by filename for reproducibility.  At most
    ``settings.pure_max_docs`` are processed.

    Returns a list of ``PureDocument`` objects with raw extracted text.
    Empty documents (extraction failed or file is blank) are skipped.
    """
    settings = get_settings()
    input_dir = settings.pure_input_path

    if not input_dir.exists():
        raise FileNotFoundError(
            f"PURE input directory not found: '{input_dir}'. "
            "Create it and place SRS documents inside."
        )

    candidates = sorted(
        p for p in input_dir.iterdir()
        if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES
    )

    if not candidates:
        logger.warning("pure_loader.no_files", dir=str(input_dir))
        return []

    docs: list[PureDocument] = []
    for path in candidates:
        if len(docs) >= settings.pure_max_docs:
            break

        suffix = path.suffix.lower()
        extractor = _EXTRACTORS.get(suffix)
        if extractor is None:
            continue

        try:
            text = extractor(path)
        except Exception as exc:
            logger.error(
                "pure_loader.extraction_failed",
                path=str(path),
                error=str(exc)[:200],
            )
            continue

        if not text.strip():
            logger.warning("pure_loader.empty_doc", path=str(path))
            continue

        doc = PureDocument(
            doc_id=path.stem,
            path=path,
            file_type=suffix,
            text=text,
        )
        docs.append(doc)
        logger.info(
            "pure_loader.doc_loaded",
            doc_id=doc.doc_id,
            file_type=suffix,
            chars=doc.char_count,
        )

    logger.info("pure_loader.done", total_docs=len(docs))
    return docs
